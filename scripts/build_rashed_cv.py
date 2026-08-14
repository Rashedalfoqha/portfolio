from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\rashe\Downloads\Rashed_Mohammad_Alfuqaha_ATS_CV_2026.docx")
NAVY = RGBColor(26, 54, 93)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(75, 75, 75)
BLUE = RGBColor(30, 92, 160)
FONT = "Arial"


def set_font(run, size=9.2, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_free_paragraph(p, before=0, after=0, line=1.0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = True


def add_hyperlink(paragraph, text, url, color=BLUE, underline=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    r_pr.append(color_el)
    if underline:
        underline_el = OxmlElement("w:u")
        underline_el.set(qn("w:val"), "single")
        r_pr.append(underline_el)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(size)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color="1A365D", size="8", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def section_heading(doc, text):
    p = doc.add_paragraph()
    set_cell_free_paragraph(p, before=4, after=2, line=1.0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=10.4, bold=True, color=NAVY)
    run.font.letter_spacing = Pt(0.9)
    add_bottom_border(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="CV Bullet")
    set_cell_free_paragraph(p, after=1.4, line=1.03)
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_font(run, size=8.85)
    return p


def add_role(doc, title, date, subtitle=None):
    p = doc.add_paragraph()
    set_cell_free_paragraph(p, before=1.5, after=1, line=1.0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(title)
    set_font(run, size=9.2, bold=True)
    if subtitle:
        sub = p.add_run(f" | {subtitle}")
        set_font(sub, size=8.8, color=MUTED)
    date_run = p.add_run(f"\t{date}")
    set_font(date_run, size=8.7, color=MUTED)


def add_project(doc, name, stack, bullets):
    p = doc.add_paragraph()
    set_cell_free_paragraph(p, before=1.8, after=0.6, line=1.0)
    p.paragraph_format.keep_with_next = True
    name_run = p.add_run(name)
    set_font(name_run, size=9.2, bold=True)
    stack_run = p.add_run(f" | {stack}")
    set_font(stack_run, size=8.45, italic=True, color=MUTED)
    for item in bullets:
        add_bullet(doc, item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.43)
section.bottom_margin = Inches(0.43)
section.left_margin = Inches(0.58)
section.right_margin = Inches(0.58)
section.header_distance = Inches(0.2)
section.footer_distance = Inches(0.2)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(9.2)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.03

bullet_style = doc.styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH)
bullet_style.base_style = normal
bullet_style.paragraph_format.left_indent = Inches(0.18)
bullet_style.paragraph_format.first_line_indent = Inches(-0.12)
bullet_style.paragraph_format.space_after = Pt(1.4)
bullet_style.paragraph_format.line_spacing = 1.03
bullet_style._element.get_or_add_pPr()
num_pr = OxmlElement("w:numPr")
ilvl = OxmlElement("w:ilvl")
ilvl.set(qn("w:val"), "0")
num_id = OxmlElement("w:numId")
num_id.set(qn("w:val"), "1")
num_pr.append(ilvl)
num_pr.append(num_id)
bullet_style._element.pPr.append(num_pr)

name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_paragraph(name, after=1, line=1.0)
run = name.add_run("RASHED MOHAMMAD ALFUQAHA")
set_font(run, size=19.5, bold=True, color=NAVY)

role = doc.add_paragraph()
role.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_paragraph(role, after=2, line=1.0)
run = role.add_run("Full-Stack Software Engineer")
set_font(run, size=11.2, color=MUTED)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_cell_free_paragraph(contact, after=3, line=1.0)
for idx, (label, url) in enumerate(
    [
        ("Amman, Jordan", None),
        ("+962 77 170 9080", "tel:+962771709080"),
        ("rashedmohammadalfuqaha@gmail.com", "mailto:rashedmohammadalfuqaha@gmail.com"),
        ("rashedalfouqaha.netlify.app", "https://rashedalfouqaha.netlify.app/"),
        ("linkedin.com/in/rashedalfuqaha", "https://www.linkedin.com/in/rashedalfuqaha/"),
        ("github.com/Rashedalfoqha", "https://github.com/Rashedalfoqha"),
    ]
):
    if idx:
        sep = contact.add_run("  |  ")
        set_font(sep, size=8.9, color=MUTED)
    if url:
        add_hyperlink(contact, label, url)
    else:
        plain = contact.add_run(label)
        set_font(plain, size=8.9, color=MUTED)

section_heading(doc, "Professional Summary")
p = doc.add_paragraph()
set_cell_free_paragraph(p, after=2.5, line=1.08)
summary = (
    "Full-Stack Software Engineer delivering production web applications with TypeScript, React, "
    "Next.js, Node.js, and NestJS. Builds responsive product interfaces, REST APIs, authentication "
    "and RBAC workflows, real-time features, and PostgreSQL/MongoDB data layers. Uses AI-assisted "
    "development for faster exploration and implementation while retaining ownership of architecture, "
    "code review, testing, security, and final engineering decisions."
)
set_font(p.add_run(summary), size=9.0)

section_heading(doc, "Technical Skills")
skills = [
    ("Languages", "JavaScript, TypeScript, Python (emerging)"),
    ("Front-End", "React, Next.js, Redux, SCSS, Tailwind CSS, Material UI, Responsive UI"),
    ("Back-End", "Node.js, Express.js, NestJS, REST APIs, GraphQL, Socket.IO"),
    ("Data", "PostgreSQL, MongoDB, Firebase"),
    ("Tools & Delivery", "Docker, Git, GitHub, Postman, VS Code, Figma Plugin API"),
    ("Engineering", "Authentication, RBAC, API Integration, Real-Time Systems, Package Development, Microservices (emerging)"),
    ("AI-Augmented Workflow", "Codex, Claude, Cursor, GLM - reviewed through type checks, linting, tests, and manual validation"),
]
for label, value in skills:
    p = doc.add_paragraph()
    set_cell_free_paragraph(p, after=0.7, line=1.0)
    set_font(p.add_run(f"{label}: "), size=8.8, bold=True)
    set_font(p.add_run(value), size=8.8)

section_heading(doc, "Professional Experience")
add_role(doc, "Software Engineer - GoldenTik (formerly CartBuzz), Amman, Jordan", "Jan 2026 - Present")
add_bullet(
    doc,
    "Deliver full-stack changes for a private multi-vendor e-commerce platform using Next.js, NestJS, TypeScript, SCSS, Docker, and project-required integrations.",
)
add_bullet(
    doc,
    "Implement and review responsive interfaces, REST API integrations, and data-layer changes while protecting confidential product and business information.",
)
add_bullet(
    doc,
    "Apply Codex, Claude, Cursor, and GLM to scoped exploration and implementation, then validate generated changes through diff review, type checks, linting, tests, and manual verification.",
)

add_role(doc, "Freelance Full-Stack Developer - Independent & Vero IT", "May 2025 - Dec 2025")
add_bullet(
    doc,
    "Delivered responsive full-stack web applications using React, Next.js, Node.js, PostgreSQL, and MongoDB from requirements analysis through deployment.",
)
add_bullet(
    doc,
    "Built REST APIs, authentication and RBAC flows, real-time features, and responsive product interfaces independently and within a delivery team.",
)

add_role(doc, "Freelance CAD Designer - Independent", "Jul 2022 - Oct 2023")
add_bullet(
    doc,
    "Produced precise mosaic and architectural patterns and coordinated technical specifications with artisans and production teams.",
)

section_heading(doc, "Selected Projects")
add_project(
    doc,
    "EV Solution JO",
    "React, Node.js, MongoDB, PostgreSQL",
    [
        "Built a full-stack operations platform for DC/AC EV charging with live station and session monitoring, reporting, and analytics.",
    ],
)
add_project(
    doc,
    "Booster Icon System",
    "TypeScript, Node.js, Next.js, SVG, SVGO",
    [
        "Engineered a framework-agnostic package pipeline for 25,396 SVG files with naming, security, parity, and runtime validation gates.",
    ],
)
add_project(
    doc,
    "Figma Design Intelligence Plugin",
    "TypeScript, Figma Plugin API",
    [
        "Built a plugin that extracts design-system structure, prototype graphs, Smart Animate relationships, and interaction metadata.",
    ],
)

section_heading(doc, "Education")
add_role(doc, "Immersive Full-Stack Web Development - MERAKI Academy", "Oct 2023 - Mar 2024")
p = doc.add_paragraph()
set_cell_free_paragraph(p, after=1.2, line=1.0)
set_font(p.add_run("400+ hours of hands-on, project-based full-stack development training."), size=8.75)
add_role(doc, "Bachelor's Degree in Islamic Arts - WISE University", "Jul 2019 - Jul 2023")

section_heading(doc, "Languages")
p = doc.add_paragraph()
set_cell_free_paragraph(p, after=0, line=1.0)
set_font(
    p.add_run(
        "Arabic (Native)  |  Turkish (Fluent)  |  English (Basic Conversational Proficiency)"
    ),
    size=8.8,
)

doc.core_properties.title = "Rashed Mohammad Alfuqaha - Full-Stack Software Engineer CV"
doc.core_properties.subject = "ATS-ready software engineering resume"
doc.core_properties.author = "Rashed Mohammad Alfuqaha"
doc.core_properties.keywords = (
    "Full-Stack Software Engineer, TypeScript, React, Next.js, Node.js, NestJS, "
    "PostgreSQL, MongoDB, REST APIs, Socket.IO, Docker"
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
