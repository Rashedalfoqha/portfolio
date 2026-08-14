from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
PROFILE = json.loads((ROOT / "config" / "profile.json").read_text(encoding="utf-8"))
SETTINGS = json.loads((ROOT / "config" / "settings.json").read_text(encoding="utf-8"))
JOBS = json.loads((ROOT / "state" / "jobs.json").read_text(encoding="utf-8"))
OUT = ROOT / "output" / "applications"
REVIEW_OUT = ROOT / "output" / "manual-review"

INK = "1B2430"
MUTED = "536174"
ACCENT = "1F4E79"
LINE = "CAD3DD"


def slug(value: str) -> str:
    result = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return result[:70] or "role"


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def set_font(run, name="Arial", size=9, bold=False, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, **font):
    run = paragraph.add_run(text)
    set_font(run, **font)
    return run


def style_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(0.43)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.8)
    styles["Normal"].paragraph_format.space_after = Pt(2.2)
    styles["Normal"].paragraph_format.line_spacing = 1.04
    bullet_style = styles["List Bullet"]
    bullet_style.font.name = "Arial"
    bullet_style.font.size = Pt(8.55)
    bullet_style.paragraph_format.left_indent = Inches(0.18)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.12)
    bullet_style.paragraph_format.space_after = Pt(1.1)
    bullet_style.paragraph_format.line_spacing = 1.02


def add_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, text.upper(), size=10.2, bold=True, color=ACCENT)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    add_text(paragraph, text, size=8.55)
    return paragraph


def add_header(document: Document, role: str):
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(0)
    add_text(name, PROFILE["name"].upper(), size=18.5, bold=True)

    role_line = document.add_paragraph()
    role_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_line.paragraph_format.space_after = Pt(1)
    add_text(role_line, role.upper(), size=10.5, bold=True, color=ACCENT)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0.5)
    add_text(
        contact,
        f"{PROFILE['location']} | {PROFILE['phone']} | {PROFILE['email']}",
        size=8.3,
        color=MUTED,
    )

    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(4)
    add_text(
        links,
        "rashedalfouqaha.netlify.app | linkedin.com/in/rashedalfuqaha | github.com/Rashedalfoqha",
        size=8.3,
        color=ACCENT,
    )


def summary_for(job):
    skills = job.get("matched_skills") or []
    highlighted = ", ".join(skills[:6]) or "JavaScript, TypeScript, React, Next.js, Node.js, and NestJS"
    return (
        f"Full-Stack Software Engineer building maintainable web products across frontend, backend, APIs, "
        f"and data layers. Relevant strengths for the {job['title']} role include {highlighted}. "
        "Combines product-focused engineering with a design-trained eye and an AI-augmented workflow "
        "that keeps human review, type safety, testing, and maintainability at the center."
    )


def make_cv(job, folder: Path):
    document = Document()
    style_document(document)
    add_header(document, job["title"])
    add_heading(document, "Profile")
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_text(p, summary_for(job), size=8.7)

    add_heading(document, "Relevant Technical Skills")
    primary = PROFILE["skills"]["primary"]
    matched = job.get("matched_skills") or []
    ordered = list(dict.fromkeys(matched + primary + PROFILE["skills"]["secondary"]))
    skill_lines = [
        ("Core", ordered[:12]),
        ("Additional", ordered[12:22] + PROFILE["skills"]["emerging"]),
        ("AI-Augmented Workflow", PROFILE["skills"]["aiWorkflow"]),
    ]
    for label, values in skill_lines:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, f"{label}: ", size=8.4, bold=True)
        add_text(p, ", ".join(values), size=8.4)

    add_heading(document, "Experience")
    p = document.add_paragraph()
    add_text(p, "Software Engineer | GoldenTik (formerly CartBuzz)", size=8.9, bold=True)
    add_text(p, " | Jan 2026 - Present | Amman, Jordan", size=8.5, color=MUTED)
    add_bullet(
        document,
        "Contribute across a private multi-vendor e-commerce product using Next.js, NestJS, TypeScript, SCSS, Docker, and Python where required.",
    )
    add_bullet(
        document,
        "Work within confidentiality boundaries while supporting maintainable implementation across product layers.",
    )
    p = document.add_paragraph()
    add_text(p, "Full-Stack Developer (Contract) | Vero IT", size=8.9, bold=True)
    add_text(p, " | May 2025 - Dec 2025 | Remote", size=8.5, color=MUTED)
    add_bullet(
        document,
        "Built client web applications with React, Next.js, Node.js, Express.js, TypeScript, PostgreSQL, and MongoDB.",
    )
    add_bullet(
        document,
        "Implemented REST APIs, authentication flows, real-time features, and product-specific data models.",
    )
    p = document.add_paragraph()
    add_text(p, "Freelance CAD Designer | Self-Employed", size=8.9, bold=True)
    add_text(p, " | Jul 2022 - Oct 2023 | Amman / Remote", size=8.5, color=MUTED)
    add_bullet(
        document,
        "Translated complex geometric concepts into precise, production-ready mosaic and architectural patterns.",
    )
    add_bullet(
        document,
        "Coordinated specifications with artisans and production teams, strengthening accuracy, systems thinking, and cross-functional communication.",
    )

    add_heading(document, "Selected Projects")
    relevant = sorted(
        PROFILE["projects"],
        key=lambda project: sum(
            1 for skill in (job.get("matched_skills") or []) if skill in project["stack"]
        ),
        reverse=True,
    )
    for project in relevant[:3]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_text(p, project["name"], size=8.8, bold=True)
        add_text(
            p,
            f" | {', '.join(project['stack'])}",
            size=8.2,
            color=MUTED,
        )
        add_bullet(document, project["summary"])

    add_heading(document, "Education & Languages")
    add_bullet(document, "Immersive Full-Stack Web Development, MERAKI Academy — 400+ hands-on hours (2023-2024).")
    add_bullet(document, "Bachelor of Islamic Arts, WISE University — architecture, decorative arts, and design principles (2019-2023).")
    add_bullet(document, "Arabic (Native) · Turkish (Fluent) · English (Basic Conversational Proficiency).")

    path = folder / "Rashed_Alfuqaha_Tailored_CV.docx"
    document.save(path)
    return path


def make_letter(job, folder: Path):
    document = Document()
    style_document(document)
    add_header(document, f"Application — {job['title']}")
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    add_text(p, f"Dear {job['company']} hiring team,", size=10, bold=True)
    paragraphs = [
        (
            f"I am applying for the {job['title']} role. My work centers on building web products across "
            f"React and Next.js frontends, Node.js and NestJS services, REST APIs, and MongoDB/PostgreSQL data layers."
        ),
        (
            f"The role aligns particularly well with my experience in "
            f"{', '.join((job.get('matched_skills') or [])[:6]) or 'JavaScript, TypeScript, React, Next.js, Node.js, and REST APIs'}. "
            "I currently contribute to a multi-vendor e-commerce product while respecting strict confidentiality, "
            "and my selected work includes EV operations dashboards, real-time collaboration, social communication, "
            "job matching, and role-based course management."
        ),
        (
            "Before software, I designed architectural and mosaic patterns in CAD. That background shaped how I "
            "approach interfaces and systems: understand the structure, make each relationship intentional, and "
            "verify that the final result works at both detail and product level. I also use Codex, Claude, Cursor, "
            "and GLM as development accelerators while retaining responsibility for review, correctness, and maintainability."
        ),
        (
            f"I would welcome a conversation about how this combination could contribute to {job['company']}. "
            "Thank you for reviewing my application."
        ),
    ]
    for text in paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.12
        add_text(p, text, size=9.3)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_text(p, "Rashed Mohammad Alfuqaha", size=10, bold=True)
    p2 = document.add_paragraph()
    add_text(p2, PROFILE["email"], name="Cascadia Mono", size=8, color=MUTED)
    path = folder / "Rashed_Alfuqaha_Cover_Letter.docx"
    document.save(path)
    return path


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "Name",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=18.5,
            leading=19.5,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=3,
        ),
        "role": ParagraphStyle(
            "Role",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=11.5,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{ACCENT}"),
        ),
        "contact": ParagraphStyle(
            "Contact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{MUTED}"),
        ),
        "section": ParagraphStyle(
            "Section",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=11.5,
            textColor=colors.HexColor(f"#{ACCENT}"),
            spaceBefore=5,
            spaceAfter=3,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.65,
            leading=10.35,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=2,
        ),
        "small": ParagraphStyle(
            "Small",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.7,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=1.5,
        ),
        "letter": ParagraphStyle(
            "Letter",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=14,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=11,
        ),
    }


def pdf_header(job, styles, subtitle):
    content = [
        Paragraph(PROFILE["name"], styles["name"]),
        Paragraph(subtitle, styles["role"]),
        Paragraph(
            f"{PROFILE['location']} | {PROFILE['phone']} | {PROFILE['email']}<br/>"
            "rashedalfouqaha.netlify.app | linkedin.com/in/rashedalfuqaha | github.com/Rashedalfoqha",
        styles["contact"],
        ),
    ]
    table = Table([[content]], colWidths=[190 * mm])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def section_pdf(story, title, styles):
    story.append(Paragraph(title.upper(), styles["section"]))


def bullet_pdf(story, text, styles):
    story.append(Paragraph(f'<font color="#{ACCENT}">•</font> {text}', styles["body"]))


def make_cv_pdf(job, folder: Path):
    path = folder / "Rashed_Alfuqaha_Tailored_CV.pdf"
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=13 * mm,
        leftMargin=13 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title=f"Rashed Alfuqaha — {job['title']}",
        author=PROFILE["name"],
    )
    story = [pdf_header(job, styles, job["title"])]
    section_pdf(story, "Profile", styles)
    story.append(Paragraph(summary_for(job), styles["body"]))

    section_pdf(story, "Relevant Technical Skills", styles)
    primary = PROFILE["skills"]["primary"]
    matched = job.get("matched_skills") or []
    ordered = list(dict.fromkeys(matched + primary + PROFILE["skills"]["secondary"]))
    for label, values in [
        ("Core", ordered[:12]),
        ("Additional", ordered[12:22] + PROFILE["skills"]["emerging"]),
        ("AI-Augmented Workflow", PROFILE["skills"]["aiWorkflow"]),
    ]:
        story.append(
            Paragraph(
                f"<b>{label}:</b> {', '.join(values)}",
                styles["small"],
            )
        )

    section_pdf(story, "Experience", styles)
    story.append(
        Paragraph(
            "<b>SOFTWARE ENGINEER</b>  GoldenTik (formerly CartBuzz)  |  Jan 2026 - Present",
            styles["small"],
        )
    )
    bullet_pdf(
        story,
        "Contribute across a private multi-vendor e-commerce product using Next.js, NestJS, TypeScript, SCSS, Docker, and Python where required.",
        styles,
    )
    bullet_pdf(
        story,
        "Work within confidentiality boundaries while supporting maintainable implementation across product layers.",
        styles,
    )
    story.append(
        Paragraph(
            "<b>FULL-STACK DEVELOPER (CONTRACT)</b>  Vero IT  |  May 2025 - Dec 2025",
            styles["small"],
        )
    )
    bullet_pdf(
        story,
        "Built client web applications with React, Next.js, Node.js, Express.js, TypeScript, PostgreSQL, and MongoDB.",
        styles,
    )
    bullet_pdf(
        story,
        "Implemented REST APIs, authentication flows, real-time features, and product-specific data models.",
        styles,
    )
    story.append(
        Paragraph(
            "<b>FREELANCE CAD DESIGNER</b>  Independent  |  Jul 2022 - Oct 2023",
            styles["small"],
        )
    )
    bullet_pdf(
        story,
        "Translated complex geometric concepts into precise, production-ready mosaic and architectural patterns.",
        styles,
    )
    bullet_pdf(
        story,
        "Coordinated specifications with artisans and production teams, strengthening accuracy, systems thinking, and cross-functional communication.",
        styles,
    )

    section_pdf(story, "Selected Projects", styles)
    relevant = sorted(
        PROFILE["projects"],
        key=lambda project: sum(
            1 for skill in (job.get("matched_skills") or []) if skill in project["stack"]
        ),
        reverse=True,
    )
    for project in relevant[:3]:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f"<b>{project['name']}</b> | {', '.join(project['stack'])}",
                        styles["small"],
                    ),
                    Paragraph(project["summary"], styles["body"]),
                ]
            )
        )

    section_pdf(story, "Education & Languages", styles)
    bullet_pdf(story, "Immersive Full-Stack Web Development, MERAKI Academy — 400+ hands-on hours (2023-2024).", styles)
    bullet_pdf(story, "Bachelor of Islamic Arts, WISE University — architecture, decorative arts, and design principles (2019-2023).", styles)
    bullet_pdf(story, "Arabic (Native) · Turkish (Fluent) · English (Basic Conversational Proficiency).", styles)
    doc.build(story)
    return path


def make_letter_pdf(job, folder: Path):
    path = folder / "Rashed_Alfuqaha_Cover_Letter.pdf"
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"Cover Letter — {job['title']} — {job['company']}",
        author=PROFILE["name"],
    )
    story = [
        pdf_header(job, styles, f"APPLICATION — {job['title']}"),
        Spacer(1, 16 * mm),
        Paragraph(f"<b>Dear {job['company']} hiring team,</b>", styles["letter"]),
    ]
    paragraphs = [
        (
            f"I am applying for the {job['title']} role. My work centers on building web products across "
            f"React and Next.js frontends, Node.js and NestJS services, REST APIs, and MongoDB/PostgreSQL data layers."
        ),
        (
            f"The role aligns particularly well with my experience in "
            f"{', '.join((job.get('matched_skills') or [])[:6]) or 'JavaScript, TypeScript, React, Next.js, Node.js, and REST APIs'}. "
            "I currently contribute to a multi-vendor e-commerce product while respecting strict confidentiality, "
            "and my selected work includes EV operations dashboards, real-time collaboration, social communication, "
            "job matching, and role-based course management."
        ),
        (
            "Before software, I designed architectural and mosaic patterns in CAD. That background shaped how I "
            "approach interfaces and systems: understand the structure, make each relationship intentional, and "
            "verify that the final result works at both detail and product level. I also use Codex, Claude, Cursor, "
            "and GLM as development accelerators while retaining responsibility for review, correctness, and maintainability."
        ),
        (
            f"I would welcome a conversation about how this combination could contribute to {job['company']}. "
            "Thank you for reviewing my application."
        ),
    ]
    story.extend(Paragraph(text, styles["letter"]) for text in paragraphs)
    story.append(Spacer(1, 5 * mm))
    story.append(Paragraph("<b>Rashed Mohammad Alfuqaha</b>", styles["letter"]))
    story.append(Paragraph(PROFILE["email"], styles["small"]))
    doc.build(story)
    return path


def main():
    if OUT.exists():
        shutil.rmtree(OUT)
    if REVIEW_OUT.exists():
        shutil.rmtree(REVIEW_OUT)
    OUT.mkdir(parents=True, exist_ok=True)
    REVIEW_OUT.mkdir(parents=True, exist_ok=True)
    candidates = [
        job
        for job in JOBS
        if int(job["score"]) >= int(SETTINGS["documentScore"])
        and not job.get("application_warning")
    ][: int(SETTINGS["maxDocumentsPerRun"])]
    manifest = []
    for job in candidates:
        folder = OUT / f"{job['score']}-{slug(job['company'])}-{slug(job['title'])}"
        folder.mkdir(parents=True, exist_ok=True)
        cv = make_cv(job, folder)
        letter = make_letter(job, folder)
        cv_pdf = make_cv_pdf(job, folder)
        letter_pdf = make_letter_pdf(job, folder)
        info = {
            "job": job,
            "cv": str(cv),
            "coverLetter": str(letter),
            "cvPdf": str(cv_pdf),
            "coverLetterPdf": str(letter_pdf),
        }
        (folder / "application.json").write_text(
            json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        manifest.append(info)
    (OUT / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    manual = [
        job
        for job in JOBS
        if int(job["score"]) >= int(SETTINGS["documentScore"])
        and job.get("application_warning")
    ]
    for job in manual:
        brief = [
            f"# HUMAN-WRITTEN APPLICATION BRIEF — {job['company']}",
            "",
            "## Do not submit generated wording",
            "",
            job["application_warning"],
            "",
            f"Policy source: {job.get('policy_source', '')}",
            "",
            "## Role",
            "",
            f"- {job['title']}",
            f"- Location: {job['location']}",
            f"- Match score: {job['score']}",
            f"- Job: {job['url']}",
            "",
            "## Verified evidence you can use in your own words",
            "",
            f"- Matched skills: {', '.join(job.get('matched_skills') or [])}",
            "- Current role: Software Engineer contributing to a confidential multi-vendor e-commerce product.",
            "- Relevant product work: EV operations dashboards, real-time collaboration, social communication, job matching, and RBAC course management.",
            "- Distinctive background: CAD-based architectural and mosaic design, precision, systems thinking, and cross-functional coordination.",
            "- AI workflow: use tools for development acceleration while retaining human review and responsibility. Do not mention or reuse generated application wording where the employer prohibits it.",
            "",
            "## Gaps to answer honestly",
            "",
            f"- {', '.join(job.get('missing_skills') or []) or 'No explicit gap detected by the rules engine.'}",
        ]
        (REVIEW_OUT / f"{slug(job['company'])}-{slug(job['title'])}.md").write_text(
            "\n".join(brief), encoding="utf-8"
        )
    print(
        json.dumps(
            {
                "documents": len(manifest),
                "manualReviewBriefs": len(manual),
                "output": str(OUT),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
